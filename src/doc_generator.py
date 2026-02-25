import os
from datetime import datetime
from docx import Document
from src.config import OUTPUT_DIR


class DocGenerator:
    def __init__(self):
        os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    def create_document(self, text, filename=None):
        doc = Document()
        doc.add_heading("OCR识别结果", 0)
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("-" * 50)
        
        paragraphs = text.split("\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        if filename is None:
            filename = f"ocr_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        return filepath
